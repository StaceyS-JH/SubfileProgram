from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # H1 heading
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        
        # H2 heading
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        
        # H3 heading
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        
        # Table detection (starts with |)
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table - skip separator line (contains ---)
            for tl in table_lines:
                if '---' in tl:
                    continue
                # Split by | and clean up
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                doc.add_paragraph(' | '.join(cells))
            doc.add_paragraph('')  # Add spacing after table
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Regular paragraph
        doc.add_paragraph(line)
        i += 1
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == "__main__":
    md_path = "c:/Users/staceys/OneDrive - Jack Henry/Documents/CodeGenerator/SubfileProgram/docs/HB2360-analysis.md"
    docx_path = "c:/Users/staceys/OneDrive - Jack Henry/Documents/CodeGenerator/SubfileProgram/docs/HB2360-analysis.docx"
    convert_md_to_docx(md_path, docx_path)
